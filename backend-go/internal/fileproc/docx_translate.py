#!/usr/bin/env python3
"""docx_translate.py — PDF→DOCX→翻译→DOCX→PDF（保留排版/图片；图片内文字OCR翻译）

子命令：
  extract <pdf> <cache_docx>          转换一次并输出段落文本JSON（供LLM翻译的键，与替换目标完全一致）
  apply   <cache_docx> <out_pdf> <lang>  从stdin读translations JSON，在缓存DOCX副本上替换（含图片OCR）后转PDF
  legacy  <in_path> <out_path> <lang>    单阶段直通模式（回退用）
"""
import sys, os, io, json, re, subprocess, tempfile, shutil
from pathlib import Path

TESS_LANG_MAP = {
    "zh": "chi_sim+chi_tra", "en": "eng", "ja": "jpn", "ko": "kor",
    "fr": "fra", "de": "deu", "es": "spa", "ru": "rus", "ar": "ara",
    "pt": "por", "it": "ita", "nl": "nld", "pl": "pol", "tr": "tur",
    "th": "tha", "vi": "vie", "hi": "hin",
}

def tess_lang(code: str) -> str:
    return TESS_LANG_MAP.get((code or "")[:2].lower(), "eng")

def run(cmd: list, timeout=180):
    subprocess.run(cmd, check=True, timeout=timeout, capture_output=True)

def pdf_to_docx(pdf_path: str, docx_path: str):
    import contextlib
    with contextlib.redirect_stdout(io.StringIO()):  # 含首次import pymupdf的弃用告警与转换进度，保stdout纯净JSON
        from pdf2docx import Converter
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()

def docx_to_pdf(docx_path: str, pdf_path: str):
    run(["libreoffice", "--headless", "--convert-to", "pdf",
         "--outdir", str(Path(pdf_path).parent), docx_path])
    expected = str(Path(pdf_path).parent / (Path(docx_path).stem + ".pdf"))
    if os.path.exists(expected):
        os.replace(expected, pdf_path)

# ---------- 文本归一化与查找表 ----------
_ZW_RE = re.compile(r"[\s\u200b\u200c\u200d\u2060\ufeff]+")

def _norm(s: str) -> str:
    return _ZW_RE.sub("", s or "")

def _build_lookup(translations: dict):
    """返回 (full_lookup, partial_lookup)：短键只允许整段命中，防误吞"""
    full_lk, part_lk = {}, {}
    for orig, trans in translations.items():
        k = _norm(orig)
        if not k:
            continue
        if len(k) >= 4:
            part_lk[k] = trans
        full_lk[k] = trans
    sort_desc = lambda d: dict(sorted(d.items(), key=lambda kv: -len(kv[0])))
    return sort_desc(full_lk), sort_desc(part_lk)

# ---------- w:t 级替换（绝不触碰 drawing，图片不丢） ----------
W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

def _para_t_nodes(para):
    return para._p.findall('.//' + W_NS + 't')

def _replace_para_text(para, lookup_full: dict, lookup_part: dict) -> bool:
    ts = _para_t_nodes(para)
    if not ts:
        return False
    full = "".join(t.text or "" for t in ts)
    if not full.strip():
        return False
    nf = _norm(full)

    spans = []
    trans = lookup_full.get(nf)          # A 整段命中（允许短键）
    if trans is not None:
        spans.append((0, len(full), trans))
    else:                                 # B 子串命中（仅长键）
        idx_map = [i for i, ch in enumerate(full) if not _ZW_RE.match(ch)]
        if len(idx_map) == len(nf):
            for k, t in lookup_part.items():
                start = 0
                while True:
                    i = nf.find(k, start)
                    if i < 0:
                        break
                    spans.append((idx_map[i], idx_map[i + len(k) - 1] + 1, t))
                    start = i + len(k)
    if not spans:
        return False

    spans.sort(key=lambda x: x[0])
    merged = []
    for s_, e_, t_ in spans:
        if merged and s_ < merged[-1][1]:
            continue
        merged.append((s_, e_, t_))

    out = []
    off = 0
    si = 0
    for t in ts:
        txt = t.text or ""
        L = len(txt)
        s0, s1 = off, off + L
        buf = []
        cur = s0
        j = si
        while j < len(merged) and merged[j][1] <= cur:
            j += 1
        si = j
        while j < len(merged) and merged[j][0] < s1:
            ms, me, mt = merged[j]
            a = max(ms, cur)
            b = min(me, s1)
            if a > cur:
                buf.append(txt[cur - s0:a - s0])
            if a == ms:
                buf.append(mt)
            cur = b
            j += 1
        if cur < s1:
            buf.append(txt[cur - s0:])
        out.append("".join(buf))
        off = s1

    changed = False
    for t, nt in zip(ts, out):
        if (t.text or "") != nt:
            t.text = nt
            t.set(XML_SPACE, 'preserve')
            changed = True
    return changed

def apply_translations_to_text(text: str, translations: dict) -> str:
    for orig, trans in translations.items():
        if orig and orig in text:
            text = text.replace(orig, trans)
    return text

def iter_all_paragraphs(doc):
    """正文 XML 内全部段落（含表格/嵌套表格/文本框 txbx）+ 页眉页脚。
    注意：lxml 元素代理会被回收且 id() 复用，严禁按 id 去重。"""
    from docx.text.paragraph import Paragraph
    for p_el in doc.element.body.iter(W_NS + 'p'):
        yield Paragraph(p_el, doc)
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            try:
                el = part._element
            except Exception:
                continue
            if el is None:
                continue
            for p_el in el.iter(W_NS + 'p'):
                yield Paragraph(p_el, doc)

def _has_drawing(r) -> bool:
    """run 内是否含图片/图形（w:drawing 或 w:pict）"""
    return (r._r.find('.//' + W_NS + 'drawing') is not None
            or r._r.find('.//' + W_NS + 'pict') is not None)

def translate_docx_text(docx_path: str, translations: dict):
    from docx import Document
    doc = Document(docx_path)
    lk_full, lk_part = _build_lookup(translations)
    for para in iter_all_paragraphs(doc):
        if _replace_para_text(para, lk_full, lk_part):
            continue
        # 兜底：逐 run 子串替换。★ 含图 run 必须跳过（run.text= 会删掉 drawing）
        for run in para.runs:
            if _has_drawing(run):
                continue
            new = apply_translations_to_text(run.text, translations)
            if new != run.text:
                run.text = new
    doc.save(docx_path)

# ---------- 字体归一化 ----------
# OCR 识别语言：源文档可能中英混排，固定多语组合；提取与应用两侧必须一致
OCR_LANGS = "chi_sim+chi_tra+eng"

_FONT_CACHE = None

def resolve_cjk_font() -> str:
    """解析输出用 CJK 字体（结果缓存）：
    ① 环境变量 CJK_FONT_NAME 显式指定；
    ② 候选链按序取 fontconfig 已安装且覆盖汉字的第一个：
       阿里巴巴普惠体(Alibaba PuHuiTi，默认) → 苹方 → Noto Sans CJK SC → 任一 :lang=zh；
    ③ 兜底 Noto Sans CJK SC。
    只输出「确实已安装」的字族名，杜绝 LibreOffice 回退到无汉字字体。"""
    global _FONT_CACHE
    if _FONT_CACHE:
        return _FONT_CACHE
    import os as _os, subprocess as _sp
    name = _os.environ.get("CJK_FONT_NAME", "").strip()
    if name and _font_installed(name):
        _FONT_CACHE = name
        return name
    for cand in ("Alibaba PuHuiTi", "PingFang SC", "Noto Sans CJK SC", "Noto Sans CJK JP"):
        if _font_installed(cand):
            _FONT_CACHE = cand
            return cand
    try:
        out = _sp.run(["fc-list", ":lang=zh", "family"], capture_output=True, text=True, timeout=10)
        for ln in out.stdout.splitlines():
            fam = ln.split(",")[0].strip()
            if fam:
                _FONT_CACHE = fam
                return fam
    except Exception:
        pass
    _FONT_CACHE = "Noto Sans CJK SC"
    return _FONT_CACHE

def _font_installed(family: str) -> bool:
    import subprocess as _sp
    try:
        r = _sp.run(["fc-list", family, "family"], capture_output=True, text=True, timeout=10)
        return family.lower() in (r.stdout or "").lower()
    except Exception:
        return False

def normalize_fonts(docx_path: str):
    """还原文件时的字体兜底规则：
    原文声明的字族若服务器未安装且文件未内嵌该字体（pdf2docx 重建的 DOCX
    一律不携带内嵌字体，故判定条件即「未安装」），则统一改写为 resolve_cjk_font()
    解析出的默认 CJK 字体（苹方优先，其次 Noto Sans CJK SC）。
    已安装的字体原样保留；确保任何汉字都不会落入无字形回退路径（方框）。"""
    from docx import Document
    font = resolve_cjk_font()
    doc = Document(docx_path)
    def rewrite(el):
        n = 0
        for rf in el.iter(W_NS + 'rFonts'):
            for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
                q = W_NS + attr
                cur = rf.get(q)
                if not cur:
                    continue
                if _font_installed(cur):
                    continue  # 已安装：保留原文体
                rf.set(q, font)
                n += 1
        return n
    n = rewrite(doc.element.body)
    try:
        n += rewrite(doc.styles.element)
    except Exception:
        pass
    doc.save(docx_path)

def _pil_font(size: int):
    """为 PIL 绘制解析 CJK 字体文件路径（跟随 resolve_cjk_font 的结果）。"""
    from PIL import ImageFont
    import subprocess as _sp
    fam = resolve_cjk_font()
    try:
        r = _sp.run(["fc-match", "-f", "%{file}", fam],
                    capture_output=True, text=True, timeout=10)
        path = (r.stdout or "").strip()
        if path and os.path.exists(path):
            return ImageFont.truetype(path, size)
    except Exception:
        pass
    for cand in ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",):
        if os.path.exists(cand):
            return ImageFont.truetype(cand, size)
    return ImageFont.load_default()

# ---------- 图片 OCR 翻译 ----------
def _ocr_lines(img, lang: str):
    """OCR 并按行分组：返回 [(text, x, y, w, h)]，行文本=词序拼接"""
    import pytesseract
    data = pytesseract.image_to_data(img, lang=OCR_LANGS,
                                     output_type=pytesseract.Output.DICT)
    groups = {}
    for i in range(len(data["text"])):
        w = (data["text"][i] or "").strip()
        if not w:
            continue
        key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
        groups.setdefault(key, []).append(
            (data["left"][i], data["top"][i], data["width"][i], data["height"][i], w))
    lines = []
    for key in sorted(groups):
        ws = sorted(groups[key], key=lambda t: t[0])
        text = " ".join(t[4] for t in ws)
        x0 = min(t[0] for t in ws); y0 = min(t[1] for t in ws)
        x1 = max(t[0]+t[2] for t in ws); y1 = max(t[1]+t[3] for t in ws)
        lines.append((text, x0, y0, x1-x0, y1-y0))
    return lines

def ocr_and_translate_image(img_bytes: bytes, translations: dict, lang: str) -> bytes | None:
    from PIL import Image, ImageDraw, ImageFont
    img = Image.open(io.BytesIO(img_bytes))
    if img.mode != "RGB":
        img = img.convert("RGB")
    w, h = img.size
    if w < 40 or h < 20:
        return None
    lk = { _norm(k): v for k, v in translations.items() }
    draw = ImageDraw.Draw(img)
    hit = False
    for text, x, y, bw, bh in _ocr_lines(img, lang):
        trans = None
        nk = _norm(text)
        if nk in lk:
            trans = lk[nk]
        else:
            for k, v in lk.items():
                if len(k) >= 4 and k in nk:
                    trans = v  # 行内子串命中→整行替换
                    break
        if not trans or trans == text:
            continue
        hit = True
        fs = max(8, int(bh * 0.72))
        font = _pil_font(fs)
        draw.rectangle([x, y, x + bw, y + bh], fill="white")
        draw.text((x, y), trans[:80], fill=(20, 20, 20), font=font)
    return buf_ret(img) if hit else None

def buf_ret(img):
    buf = io.BytesIO(); img.save(buf, format="PNG"); return buf.getvalue()

def translate_docx_images(docx_path: str, translations: dict, lang: str):
    from docx import Document
    try:
        doc = Document(docx_path)
    except Exception:
        return
    n = 0
    for rel in list(doc.part.rels.values()):
        if "image" not in rel.reltype:
            continue
        blob = rel.target_part.blob
        if len(blob) < 3000:  # 忽略图标/装饰小图
            continue
        try:
            nb = ocr_and_translate_image(blob, translations, lang)
            if nb:
                rel.target_part._blob = nb
                n += 1
        except Exception:
            continue
    if n:
        try:
            doc.save(docx_path)
        except Exception:
            pass


# ---------- 整页 OCR 模式（图形化/扫描版 PDF 兜底） ----------
PAGE_TEXT_MIN = 200  # 平均每页文本层字符低于此值视为图形化文档

def page_is_graphical(pdf_path: str) -> bool:
    try:
        import pymupdf
        d = pymupdf.open(pdf_path)
        n = max(1, len(d))
        chars = len("".join(p.get_text() for p in d).strip())
        return chars / n < PAGE_TEXT_MIN
    except Exception:
        return False

def pageocr_collect(pdf_path: str):
    """整页栅格化并 OCR 行键收集。返回 (keys, [(page_idx, img_bytes, lines)])"""
    import pymupdf
    from PIL import Image as PILImage
    import io as _io
    d = pymupdf.open(pdf_path)
    keys, seen, pages = [], set(), []
    for pi in range(len(d)):
        pix = d[pi].get_pixmap(matrix=pymupdf.Matrix(2, 2))  # ~144dpi
        img = PILImage.open(_io.BytesIO(pix.tobytes("png"))).convert("RGB")
        buf = _io.BytesIO(); img.save(buf, format="PNG")
        b = buf.getvalue()
        lines = _ocr_lines(img, OCR_LANGS)
        pages.append((pi, b, lines))
        for text, *_ in lines:
            k = _norm(text)
            if len(k) >= 2 and k not in seen:
                seen.add(k); keys.append(text)
    return keys, pages

def cmd_pageocr_apply(pdf_path: str, out_path: str, lang: str, translations: dict):
    """整页模式回写：命中的行白底覆盖写译文，页图重组为 PDF（版式像素级保真）。"""
    import pymupdf
    from PIL import Image, ImageDraw
    import io as _io
    lk = {}
    for k, v in translations.items():
        nk = _norm(k)
        if len(nk) >= 2:
            lk[nk] = v
    d = pymupdf.open(pdf_path)
    out = pymupdf.open()
    total_hit = 0
    for pi in range(len(d)):
        pix = d[pi].get_pixmap(matrix=pymupdf.Matrix(2, 2))
        img = Image.open(_io.BytesIO(pix.tobytes("png"))).convert("RGB")
        draw = ImageDraw.Draw(img)
        for text, x, y, bw, bh in _ocr_lines(img, OCR_LANGS):
            nk = _norm(text)
            trans = lk.get(nk)
            if not trans and len(nk) >= 6:
                for k, v in lk.items():
                    if len(k) >= 6 and k in nk:
                        trans = v  # 兜底：行内包含某长键（OCR 断行差异）
                        break
            if not trans:
                continue
            fs = max(9, int(bh * 0.72))
            font = _pil_font(fs)
            draw.rectangle([x - 1, y - 1, x + bw + 1, y + bh + 1], fill="white")
            draw.text((x, y), trans[:120], fill=(20, 20, 20), font=font)
            total_hit += 1
        w, h = img.size
        pg = out.new_page(width=w * 72 / 144, height=h * 72 / 144)
        pg.insert_image(pg.rect, stream=_io.BytesIO(buf_ret(img)))
        d.close() if False else None
    out.save(out_path)
    print(f"OK pageocr hits={total_hit}: {out_path}")


# ---------- 表格自适应 ----------
def normalize_tables(docx_path: str):
    """所有表格：layout=autofit（列宽随内容自适应）+ 总宽100%。
    译文与原文长度差异大，固定列宽会导致溢出/换行崩坏。"""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    doc = Document(docx_path)
    n = 0
    for tbl in doc.element.body.iter(qn('w:tbl')):
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        layout = tblPr.find(qn('w:tblLayout'))
        if layout is None:
            layout = OxmlElement('w:tblLayout')
            tblPr.append(layout)
        layout.set(qn('w:type'), 'autofit')
        tw = tblPr.find(qn('w:tblW'))
        if tw is None:
            tw = OxmlElement('w:tblW')
            tblPr.append(tw)
        tw.set(qn('w:w'), '5000')
        tw.set(qn('w:type'), 'pct')  # 100%
        n += 1
    doc.save(docx_path)
# ---------- extract / apply ----------
def cmd_extract(pdf_path: str, cache_docx: str):
    # 图形化文档（文本层稀薄）：跳过 docx 重建，直接整页 OCR 收集行键
    if page_is_graphical(pdf_path):
        keys, _pages = pageocr_collect(pdf_path)
        print(json.dumps({"success": True, "mode": "pageocr", "texts": keys}, ensure_ascii=False))
        return
    pdf_to_docx(pdf_path, cache_docx)
    from docx import Document
    doc = Document(cache_docx)
    seen, texts = set(), []
    for p in iter_all_paragraphs(doc):
        txt = "".join(t.text or "" for t in _para_t_nodes(p)).strip()
        if not txt:
            continue
        k = _norm(txt)
        if len(k) < 2 or k in seen:
            continue
        seen.add(k)
        texts.append(txt)
    # ★ 图片 OCR 行键：让 LLM 连图内文字一并翻译（apply 阶段按行精确命中）
    try:
        from docx import Document as _Doc
        from PIL import Image as _Image
        import io as _io, pytesseract as _pt
        d2 = _Doc(cache_docx)
        for rel in d2.part.rels.values():
            if "image" not in rel.reltype:
                continue
            blob = rel.target_part.blob
            if len(blob) < 3000:
                continue
            try:
                im = _Image.open(_io.BytesIO(blob))
                if im.mode != "RGB":
                    im = im.convert("RGB")
                if im.size[0] < 40 or im.size[1] < 20:
                    continue
                for ln_text, *_ in _ocr_lines(im, OCR_LANGS):
                    k2 = _norm(ln_text)
                    if len(k2) >= 2 and k2 not in seen:
                        seen.add(k2)
                        texts.append(ln_text)
            except Exception:
                continue
    except Exception:
        pass
    print(json.dumps({"success": True, "mode": "docx", "texts": texts}, ensure_ascii=False))

def cmd_apply(cache_docx: str, out_path: str, lang: str, translations: dict):
    tmp = tempfile.mkdtemp()
    try:
        work = os.path.join(tmp, "work.docx")
        shutil.copy2(cache_docx, work)
        if translations:
            translate_docx_text(work, translations)
            normalize_fonts(work)   # 字体兜底：未装字族→默认CJK（普惠体），汉字不回退
            normalize_tables(work)  # ★ 表格自适应：列宽自动布局+表宽100%，防译文溢出错位
            # 图片内嵌 OCR 翻译暂时停用（产品决策 2026-08-25）：避免半译状态破坏观感
        docx_to_pdf(work, out_path)
        print(f"OK: {out_path}")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

def cmd_legacy(in_path: str, out_path: str, lang: str, translations: dict):
    ext = Path(in_path).suffix.lower()
    tmp = tempfile.mkdtemp()
    try:
        if ext == ".pdf":
            cache = os.path.join(tmp, "in.docx")
            pdf_to_docx(in_path, cache)
        elif ext == ".docx":
            cache = in_path
        else:
            print(f"Unsupported format: {ext}", file=sys.stderr)
            sys.exit(1)
        cmd_apply(cache if ext == ".pdf" else cache, out_path, lang, translations)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

def main():
    argv = sys.argv[1:]
    if not argv:
        print(__doc__, file=sys.stderr); sys.exit(1)
    mode = argv[0]
    if mode == "extract" and len(argv) >= 3:
        cmd_extract(argv[1], argv[2]); return
    if mode == "pageocr" and len(argv) >= 4:
        raw = sys.stdin.read()
        try:
            data = json.loads(raw)
        except Exception:
            data = {}
        cmd_pageocr_apply(argv[1], argv[2], argv[3], data.get("translations", {})); return
    if mode == "apply" and len(argv) >= 4:
        raw = sys.stdin.read()
        try:
            data = json.loads(raw)
        except Exception:
            data = {}
        cmd_apply(argv[1], argv[2], argv[3], data.get("translations", {})); return
    if mode == "legacy" and len(argv) >= 4:
        raw = sys.stdin.read()
        try:
            data = json.loads(raw)
        except Exception:
            data = {}
        cmd_legacy(argv[1], argv[2], argv[3], data.get("translations", {})); return
    # 兼容旧调用：<in> <out> <lang>
    if len(argv) >= 3:
        raw = sys.stdin.read()
        try:
            data = json.loads(raw)
        except Exception:
            data = {}
        cmd_legacy(argv[0], argv[1], argv[2], data.get("translations", {})); return
    print(__doc__, file=sys.stderr); sys.exit(1)

if __name__ == "__main__":
    main()
