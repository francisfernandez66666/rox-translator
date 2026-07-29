# ============================================================================
# services/file_service.py — 通用文件保存服务
# ============================================================================
# 【作用】统一管理所有技能的文件输出，确保文件保存到规范目录
# 【用法】
#   from services.file_service import save_text_file, save_binary_file
#   path = save_text_file("PRD内容...", "极石车机语音助手", ".docx")
#   path = save_binary_file(binary_data, "图片", ".png")
# ============================================================================

import os
import time
from pathlib import Path

# ---- 输出目录 ----
# 优先使用用户数据目录（避免 App Translocation 只读问题）
_USER_DATA_DIR = os.getenv("USER_DATA_DIR")
if _USER_DATA_DIR:
    # 如果设置了用户数据目录环境变量，输出目录放在其 _output 子目录下
    OUTPUT_DIR = os.path.join(_USER_DATA_DIR, "_output")
else:
    # 默认输出到用户 Library/Application Support 下的 _output 目录
    OUTPUT_DIR = os.path.join(str(Path.home()), "Library", "Application Support", "翻译助手", "_output")
# 确保输出目录存在，不存在则创建
os.makedirs(OUTPUT_DIR, exist_ok=True)


def _safe_filename(name: str, max_len: int = 30) -> str:
    """
    生成安全的文件名（去除特殊字符，限制长度）

    Args:
        name: 原始文件名（可能含特殊字符）
        max_len: 最大长度，默认30字符

    Returns:
        处理后的安全文件名（不含扩展名）
    """
    # 只保留字母、数字、下划线、点、短横线和中文全角空格
    safe = "".join(c for c in name if c.isalnum() or c in "._- 　").strip()
    # 将普通空格替换为下划线，并截断到最大长度
    safe = safe.replace(" ", "_")[:max_len]
    # 如果结果为空，使用时间戳作为文件名
    if not safe:
        safe = f"output_{int(time.time())}"
    return safe


def save_text_file(content: str, name: str, ext: str = ".md", subdir: str = "") -> str:
    """
    保存文本内容到文件

    Args:
        content: 文件内容
        name: 文件名（不含扩展名，会自动清理）
        ext: 扩展名（.md / .docx / .txt 等）
        subdir: 子目录名（可选，如 "prd" / "translation"）

    Returns:
        保存的文件绝对路径
    """
    # 确定输出目录：如果指定了子目录，则在其下创建
    out_dir = os.path.join(OUTPUT_DIR, subdir) if subdir else OUTPUT_DIR
    os.makedirs(out_dir, exist_ok=True)

    # 生成文件名
    filename = f"{_safe_filename(name)}{ext}"
    filepath = os.path.join(out_dir, filename)

    # 如果文件已存在，通过添加序号避免覆盖（如 文件名_1.md, 文件名_2.md）
    counter = 1
    while os.path.exists(filepath):
        filename = f"{_safe_filename(name)}_{counter}{ext}"
        filepath = os.path.join(out_dir, filename)
        counter += 1

    # 写入文件（UTF-8编码）
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    print(f"  📄 [file_service] 文件已保存: {filename}")
    return filepath


def save_binary_file(data: bytes, name: str, ext: str = ".png", subdir: str = "") -> str:
    """
    保存二进制数据到文件

    Args:
        data: 二进制数据
        name: 文件名（不含扩展名）
        ext: 扩展名（.png / .jpg 等）
        subdir: 子目录名（可选）

    Returns:
        保存的文件绝对路径
    """
    # 确定输出目录
    out_dir = os.path.join(OUTPUT_DIR, subdir) if subdir else OUTPUT_DIR
    os.makedirs(out_dir, exist_ok=True)

    # 生成文件名，避免重名
    filename = f"{_safe_filename(name)}{ext}"
    filepath = os.path.join(out_dir, filename)

    counter = 1
    while os.path.exists(filepath):
        filename = f"{_safe_filename(name)}_{counter}{ext}"
        filepath = os.path.join(out_dir, filename)
        counter += 1

    # 以二进制写模式写入文件
    with open(filepath, "wb") as f:
        f.write(data)

    print(f"  📄 [file_service] 文件已保存: {filename}")
    return filepath


def save_docx(md_content: str, name: str, subdir: str = "") -> str:
    """
    将 Markdown 内容转换为 Word 文档并保存（三层保障）

    三层降级策略：
      第一层：专业转换 — 保留标题层级、列表样式、加粗等格式
      第二层：简单转换 — 只保留纯段落，无格式
      第三层：兜底保存 — 所有 docx 转换失败时，保存为 .md 纯文本

    Args:
        md_content: Markdown 格式内容
        name: 文件名（不含扩展名）
        subdir: 子目录名（可选）

    Returns:
        保存的 .docx 文件绝对路径
    """
    out_dir = os.path.join(OUTPUT_DIR, subdir) if subdir else OUTPUT_DIR
    os.makedirs(out_dir, exist_ok=True)

    filename = f"{_safe_filename(name)}.docx"
    filepath = os.path.join(out_dir, filename)

    counter = 1
    while os.path.exists(filepath):
        filename = f"{_safe_filename(name)}_{counter}.docx"
        filepath = os.path.join(out_dir, filename)
        counter += 1

    # ---- 第一层：专业 docx 转换（保留格式） ----
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 设置默认字体为微软雅黑，大小11磅
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(11)

        # 逐行解析 Markdown 并转换为 Word 段落
        for line in md_content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue

            # 标题层级（# ~ ####）
            if stripped.startswith('# '):
                p = doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('## '):
                p = doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('### '):
                p = doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('#### '):
                p = doc.add_heading(stripped[5:], level=4)
            # 列表项（无序列表 - 或 *）
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            # 列表项（有序列表 1. 2. 等）
            elif len(stripped) > 2 and stripped[0].isdigit() and '. ' in stripped[:5]:
                doc.add_paragraph(stripped.split('. ', 1)[1], style='List Number')
            else:
                # 普通段落（简单处理 **加粗** 格式）
                p = doc.add_paragraph()
                # 按 ** 分割，奇数索引段加粗
                parts = stripped.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        run = p.add_run(part)
                    else:
                        run = p.add_run(part)
                        run.bold = True

        doc.save(filepath)
        print(f"  📄 [file_service] Word 文档已保存: {filename}")
        return filepath

    except Exception as e1:
        print(f"  ⚠️ [file_service] 专业 docx 转换失败: {e1}")

    # ---- 第二层：简单 docx（纯段落） ----
    try:
        from docx import Document
        doc = Document()
        for line in md_content.split('\n'):
            stripped = line.strip()
            if stripped:
                doc.add_paragraph(stripped)
        doc.save(filepath)
        print(f"  📄 [file_service] 简单 Word 文档已保存: {filename}")
        return filepath

    except Exception as e2:
        print(f"  ⚠️ [file_service] 简单 docx 也失败: {e2}")

    # ---- 第三层：兜底保存为 .md ----
    md_filepath = filepath.replace('.docx', '.md')
    with open(md_filepath, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"  📄 [file_service] docx 全部失败，保存为 Markdown: {md_filepath}")
    return md_filepath
